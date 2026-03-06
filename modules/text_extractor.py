import os
from io import BytesIO
import PyPDF2
from docx import Document


def extract_txt(file):
    if isinstance(file, str):
        with open(file, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    else:
        return file.read().decode("utf-8", errors="ignore")


def extract_pdf(file):
    text = ""

    if isinstance(file, str):
        with open(file, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() + "\n"
    else:
        reader = PyPDF2.PdfReader(BytesIO(file.read()))
        for page in reader.pages:
            text += page.extract_text() + "\n"

    return text


def extract_docx(file):
    text = ""

    if isinstance(file, str):
        doc = Document(file)
    else:
        doc = Document(BytesIO(file.read()))

    for para in doc.paragraphs:
        text += para.text + "\n"

    return text


def extract_text(file):

    if isinstance(file, str):
        extension = os.path.splitext(file)[1].lower()
    else:
        extension = os.path.splitext(file.name)[1].lower()

    if extension == ".txt":
        return extract_txt(file)

    elif extension == ".pdf":
        return extract_pdf(file)

    elif extension == ".docx":
        return extract_docx(file)

    else:
        raise ValueError("Unsupported file format")